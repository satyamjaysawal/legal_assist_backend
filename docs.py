import base64
import io
import os
import re
from pathlib import Path
from typing import Any

from dotenv import load_dotenv
from groq import Groq

load_dotenv(Path(__file__).resolve().parent / ".env")
load_dotenv(Path(__file__).resolve().parent.parent / "legal_assist" / ".env")

GROQ_API_KEY = os.getenv("GROQ_API_KEY")
GROQ_VISION_MODEL = os.getenv("GROQ_VISION_MODEL", "qwen/qwen3.6-27b")
MAX_UPLOAD_BYTES = int(os.getenv("MAX_UPLOAD_BYTES", str(4 * 1024 * 1024)))
CHUNK_SIZE = int(os.getenv("DOC_CHUNK_SIZE", "900"))
CHUNK_OVERLAP = int(os.getenv("DOC_CHUNK_OVERLAP", "150"))

PDF_TYPES = {"application/pdf"}
DOCX_TYPES = {
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
}
TEXT_TYPES = {
    "text/plain",
    "text/markdown",
    "text/csv",
    "application/json",
}
IMAGE_TYPES = {
    "image/png",
    "image/jpeg",
    "image/jpg",
    "image/webp",
    "image/gif",
}

EXT_KIND = {
    ".pdf": "pdf",
    ".docx": "docx",
    ".txt": "text",
    ".md": "text",
    ".csv": "text",
    ".json": "text",
    ".png": "image",
    ".jpg": "image",
    ".jpeg": "image",
    ".webp": "image",
    ".gif": "image",
}


def detect_kind(filename: str, content_type: str) -> str:
    ext = os.path.splitext(filename or "")[1].lower()
    if ext in EXT_KIND:
        return EXT_KIND[ext]
    ctype = (content_type or "").split(";")[0].strip().lower()
    if ctype in PDF_TYPES:
        return "pdf"
    if ctype in DOCX_TYPES:
        return "docx"
    if ctype in IMAGE_TYPES:
        return "image"
    if ctype in TEXT_TYPES or ctype.startswith("text/"):
        return "text"
    raise ValueError(f"Unsupported file type: {filename or ctype or 'unknown'}")


def parse_pdf(data: bytes) -> str:
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(data))
    pages = []
    for i, page in enumerate(reader.pages, start=1):
        text = (page.extract_text() or "").strip()
        if text:
            pages.append(f"[Page {i}]\n{text}")
    return "\n\n".join(pages).strip()


def parse_docx(data: bytes) -> str:
    from docx import Document

    document = Document(io.BytesIO(data))
    parts = [para.text.strip() for para in document.paragraphs if para.text.strip()]
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts).strip()


def parse_text(data: bytes) -> str:
    for encoding in ("utf-8", "utf-16", "latin-1"):
        try:
            return data.decode(encoding).strip()
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="replace").strip()


def parse_image(data: bytes, content_type: str, filename: str) -> str:
    if not GROQ_API_KEY:
        raise RuntimeError("GROQ_API_KEY is required to read images")
    mime = (content_type or "image/png").split(";")[0].strip() or "image/png"
    if mime == "image/jpg":
        mime = "image/jpeg"
    encoded = base64.b64encode(data).decode("ascii")
    client = Groq(api_key=GROQ_API_KEY)
    result = client.chat.completions.create(
        model=GROQ_VISION_MODEL,
        temperature=0,
        max_tokens=1800,
        messages=[
            {
                "role": "user",
                "content": [
                    {
                        "type": "text",
                        "text": (
                            "Extract every readable word from this legal document image. "
                            "Keep original order. If it is a form, keep labels with values. "
                            "Return plain text only."
                        ),
                    },
                    {
                        "type": "image_url",
                        "image_url": {"url": f"data:{mime};base64,{encoded}"},
                    },
                ],
            }
        ],
    )
    text = (result.choices[0].message.content or "").strip()
    text = re.sub(r"<think>.*?</think>", "", text, flags=re.DOTALL).strip()
    if not text:
        raise RuntimeError(f"No text could be read from image {filename}")
    return text


def parse_file(filename: str, content_type: str, data: bytes) -> dict[str, Any]:
    if not data:
        raise ValueError("Empty file")
    if len(data) > MAX_UPLOAD_BYTES:
        raise ValueError(f"File is larger than {MAX_UPLOAD_BYTES // (1024 * 1024)} MB")
    kind = detect_kind(filename, content_type)
    if kind == "pdf":
        text = parse_pdf(data)
    elif kind == "docx":
        text = parse_docx(data)
    elif kind == "image":
        text = parse_image(data, content_type, filename)
    else:
        text = parse_text(data)
    text = re.sub(r"\n{3,}", "\n\n", text or "").strip()
    if not text:
        raise ValueError(f"No text extracted from {filename}")
    return {
        "filename": filename,
        "kind": kind,
        "content_type": content_type,
        "bytes": len(data),
        "text": text,
        "chars": len(text),
    }


def chunk_text(text: str) -> list[str]:
    clean = re.sub(r"[ \t]+", " ", text or "").strip()
    if not clean:
        return []
    if len(clean) <= CHUNK_SIZE:
        return [clean]
    chunks: list[str] = []
    start = 0
    while start < len(clean):
        end = min(len(clean), start + CHUNK_SIZE)
        if end < len(clean):
            cut = clean.rfind(" ", start + CHUNK_SIZE // 2, end)
            if cut > start:
                end = cut
        piece = clean[start:end].strip()
        if piece:
            chunks.append(piece)
        if end >= len(clean):
            break
        start = max(end - CHUNK_OVERLAP, start + 1)
    return chunks
