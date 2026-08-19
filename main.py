import json
import os
import uuid
from pathlib import Path

from dotenv import load_dotenv

from fastapi import Depends, FastAPI, File, Form, HTTPException, UploadFile, WebSocket
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse
from pydantic import BaseModel, Field

from auth import (
    ALL_ROLES, ROLE_GUEST, ROLE_USER, ROLE_LAWYER, ROLE_ADMIN,
    current_user, login_user, make_token, register_user, update_user,
    optional_user, require_role,
)
from graph import (
    DEFAULT_ANALYSIS,
    build_graph,
    latest_user_text,
    stream_graph,
    suggest_followups,
    suggest_title,
)
from multi_graph import build_multi_graph, stream_multi_graph
from agents.base import list_agents as list_registered_agents
from connectors import list_connectors
from connectors.base import get_connector
from websocket.lawyer_connect import (
    create_room, get_room, list_rooms, close_room,
    handle_user_websocket, handle_lawyer_websocket,
)
from journeys import create_journey, delete_all_journeys, delete_journey, get_journey, list_journeys, rename_journey
from cache import get_prompt_cache, set_prompt_cache
from docs import MAX_UPLOAD_BYTES, MAX_UPLOAD_MB, chunk_text, parse_file
from embeddings import embed_status
from files import files_status, get_original_file, store_original_file
from memory import (
    compress_history,
    get_procedural_memory,
    get_user_profile,
    layer_status,
    list_episodes,
    list_user_facts,
    load_all,
    load_user_profile,
    save_all,
    update_user_profile,
)
from vectordb import (
    delete_doc,
    format_hits,
    get_doc,
    hits_fingerprint,
    ingest_document,
    list_docs,
    qdrant_status,
    search_docs,
)

ROOT = Path(__file__).resolve().parent
load_dotenv(ROOT / ".env")
load_dotenv(ROOT.parent / "legal_assist" / ".env")

GROQ_API_KEY = os.getenv("GROQ_API_KEY")
GROQ_MODEL = os.getenv("GROQ_MODEL", "openai/gpt-oss-120b")

app = FastAPI(title="Legal AI Assistant")

extra_origins = [
    origin.strip()
    for origin in os.getenv("CORS_ORIGINS", "").split(",")
    if origin.strip()
]

app.add_middleware(
    CORSMiddleware,
    allow_origins=[
        "http://localhost:5173",
        "http://127.0.0.1:5173",
        *extra_origins,
    ],
    allow_origin_regex=r"https://.*\.vercel\.app",
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


class AuthPayload(BaseModel):
    email: str
    password: str
    name: str = ""
    role: str = "user"


class ProfileUpdate(BaseModel):
    name: str


class JourneyCreate(BaseModel):
    title: str = ""


class JourneyRename(BaseModel):
    title: str


class ChatMessage(BaseModel):
    role: str = Field(pattern="^(user|assistant)$")
    content: str


class ChatRequest(BaseModel):
    messages: list[ChatMessage]
    journey_id: str = ""
    session_id: str = ""


def require_key() -> str:
    if not GROQ_API_KEY:
        raise HTTPException(status_code=500, detail="GROQ_API_KEY is not configured")
    return GROQ_API_KEY


# ── Greeting fast-path ─────────────────────────────────────────
import re as _re

_GREETING_PATTERNS = _re.compile(
    r"^(hi+|hello|hey|heyy+|howdy|sup|yo+|namaste|good\s*(morning|afternoon|evening)|"
    r"hola|what'?s?\s*up|w+sup)[\s!?.]*$",
    _re.IGNORECASE,
)

_GREETING_REPLY = (
    "Hello! I'm your Legal AI Assistant. I can help you with:\n\n"
    "- **General legal questions** — explain concepts, rights, and procedures\n"
    "- **Document drafting** — notices, agreements, letters, petitions\n"
    "- **Legal research** — case law, statutes, comparisons\n"
    "- **Email composition** — professional legal emails\n"
    "- **Finding a lawyer** — connect with specialists\n\n"
    "Just ask me anything, and I'll route you to the right specialist agent!"
)


def _is_simple_greeting(text: str) -> bool:
    """Return True if the text is just a greeting with no substantive query."""
    cleaned = text.strip()
    if len(cleaned) > 40:
        return False
    return bool(_GREETING_PATTERNS.match(cleaned))


def cleaned_messages(req: ChatRequest) -> list[dict[str, str]]:
    messages = [
        {"role": item.role, "content": item.content.strip()}
        for item in req.messages
        if item.content.strip()
    ]
    if not messages:
        raise HTTPException(status_code=400, detail="messages cannot be empty")
    return messages


def resolve_journey(user: dict, journey_id: str) -> dict:
    journey_id = (journey_id or "").strip()
    if journey_id:
        return get_journey(user["user_id"], journey_id)
    return create_journey(user["user_id"])


def sse(event: dict) -> str:
    return f"data: {json.dumps(event, ensure_ascii=False)}\n\n"


@app.get("/")
def root():
    return {"service": "legal_assist_backend", "ok": True}


@app.get("/health")
def health():
    return {
        "ok": True,
        "provider": "groq",
        "model": GROQ_MODEL,
        "configured": bool(GROQ_API_KEY),
        "stack": {
            "langchain": True,
            "langgraph": True,
            "multi_agent": True,
            "streaming": "langgraph.stream + multi_agent",
            "query_analyser": True,
            "auth": "mongodb",
            "roles": ALL_ROLES,
            "memory": ["in_memory", "short_term_redis", "user_thread_mongo", "long_term_mongodb"],
            "prompt_cache": True,
            "rag": True,
            "uploads": ["pdf", "docx", "text", "image"],
            "embed_model": os.getenv("GROQ_EMBED_MODEL", "nomic-embed-text-v1.5"),
            "file_store": "mongodb_gridfs",
            "max_upload_mb": MAX_UPLOAD_MB,
            "websocket_lawyer_connect": True,
        },
        "agents": list_registered_agents(),
        "connectors": list_connectors(),
        "memory": layer_status(),
        "qdrant": qdrant_status(),
        "embeddings": embed_status(),
        "files": files_status(),
    }


@app.post("/auth/register")
def auth_register(payload: AuthPayload):
    user = register_user(payload.email, payload.password, payload.name, payload.role)
    token = make_token(user["user_id"], user["email"], user.get("role") or ROLE_USER)
    journey = create_journey(user["user_id"], "New chat")
    return {"token": token, "user": user, "journey": journey}


@app.post("/auth/login")
def auth_login(payload: AuthPayload):
    user = login_user(payload.email, payload.password)
    token = make_token(user["user_id"], user["email"], user.get("role") or ROLE_USER)
    journeys = list_journeys(user["user_id"])
    journey = journeys[0] if journeys else create_journey(user["user_id"], "New chat")
    return {"token": token, "user": user, "journey": journey, "journeys": journeys or [journey]}


@app.get("/auth/me")
def auth_me(user: dict = Depends(current_user)):
    journeys = list_journeys(user["user_id"])
    return {"user": user, "journeys": journeys, "journey_count": len(journeys)}


@app.patch("/auth/me")
def auth_update(payload: ProfileUpdate, user: dict = Depends(current_user)):
    return {"user": update_user(user["user_id"], payload.name)}


@app.get("/journeys")
def journeys_list(user: dict = Depends(current_user)):
    return {"journeys": list_journeys(user["user_id"])}


@app.post("/journeys")
def journeys_create(payload: JourneyCreate, user: dict = Depends(current_user)):
    return create_journey(user["user_id"], payload.title)


@app.patch("/journeys/{journey_id}")
def journeys_rename(journey_id: str, payload: JourneyRename, user: dict = Depends(current_user)):
    return rename_journey(user["user_id"], journey_id, payload.title)


@app.delete("/journeys")
def journeys_delete_all(user: dict = Depends(current_user)):
    return delete_all_journeys(user["user_id"])


@app.delete("/journeys/{journey_id}")
def journeys_delete(journey_id: str, user: dict = Depends(current_user)):
    return delete_journey(user["user_id"], journey_id)


@app.get("/memory")
def memory_detail(journey_id: str = "", user: dict = Depends(current_user)):
    journey_id = (journey_id or "").strip()
    loaded = None
    if journey_id:
        loaded = load_all(user["user_id"], journey_id, [], "")
    profile_text, profile_report = load_user_profile(user["user_id"])
    return {
        "user_id": user["user_id"],
        "journey_id": journey_id,
        "stores": layer_status(),
        "layers": loaded["layers"] if loaded else [],
        "facts": list_user_facts(user["user_id"]),
        "recalled": loaded["facts"] if loaded else [],
        "thread": loaded["history"] if loaded else [],
        "documents": list_docs(user["user_id"], journey_id) if journey_id else list_docs(user["user_id"]),
        "qdrant": qdrant_status(),
        "embeddings": embed_status(),
        "files": files_status(),
        "profile": get_user_profile(user["user_id"]),
        "profile_text": profile_text,
        "procedural": get_procedural_memory(user["user_id"]),
        "episodes": list_episodes(user["user_id"], limit=15),
        "episodic_notes": loaded.get("episodic_notes", "") if loaded else "",
        "procedural_notes": loaded.get("procedural_notes", "") if loaded else "",
    }


@app.post("/documents/upload")
async def documents_upload(
    file: UploadFile = File(...),
    journey_id: str = Form(""),
    user: dict = Depends(current_user),
):
    journey = resolve_journey(user, journey_id)
    filename = file.filename or "upload"
    content_type = file.content_type or ""
    data = await file.read()
    user_id = user["user_id"]
    journey_id = journey["journey_id"]

    def generate():
        flow = []

        def step(name, status, detail=""):
            item = {"name": name, "status": status, "detail": detail}
            flow.append(item)
            return sse({"type": "flow", "steps": flow[:], "current": name})

        try:
            yield sse({"type": "thinking", "text": f"Received {filename}…"})
            yield sse(
                {
                    "type": "file",
                    "filename": filename,
                    "bytes": len(data),
                    "content_type": content_type,
                    "max_bytes": MAX_UPLOAD_BYTES,
                }
            )
            yield step("receive", "done", f"{filename} · {len(data)} bytes")

            yield sse({"type": "thinking", "text": f"Checking size (max {MAX_UPLOAD_MB} MB)…"})
            if len(data) > MAX_UPLOAD_BYTES:
                yield step("validate", "error", f"Larger than {MAX_UPLOAD_MB} MB")
                yield sse({"type": "error", "detail": f"File is larger than {MAX_UPLOAD_MB} MB"})
                return
            yield step("validate", "done", f"Within {MAX_UPLOAD_MB} MB · {content_type or 'unknown type'}")

            yield sse({"type": "thinking", "text": f"Parsing {filename}…"})
            yield step("parse", "running", "Extracting text")
            parsed = parse_file(filename, content_type, data)
            yield step("parse", "done", f"{parsed['kind']} · {parsed['chars']} characters")

            yield sse({"type": "thinking", "text": "Splitting into chunks…"})
            chunks = chunk_text(parsed["text"])
            if not chunks:
                yield step("chunk", "error", "No text chunks")
                yield sse({"type": "error", "detail": "No chunks produced from file"})
                return
            yield step("chunk", "done", f"{len(chunks)} chunk(s)")

            doc_id = str(uuid.uuid4())
            yield sse({"type": "thinking", "text": "Uploading original file to MongoDB…"})
            yield step("mongodb", "running", "GridFS write")
            file_meta = store_original_file(
                user_id=user_id,
                journey_id=journey_id,
                doc_id=doc_id,
                filename=filename,
                content_type=content_type,
                data=data,
            )
            yield step("mongodb", "done", file_meta["detail"])
            yield sse({"type": "mongo", "report": file_meta})

            yield sse({"type": "thinking", "text": "Embedding chunks…"})
            yield step("embed", "running", "nomic-embed-text-v1.5 / fallback")
            stored = ingest_document(
                user_id,
                journey_id,
                parsed,
                chunks,
                doc_id=doc_id,
                file_meta=file_meta,
            )
            yield step("embed", "done", f"{stored.get('embed_provider') or stored.get('embed_model')} · {stored['chunks']} vector(s)")
            yield step("qdrant", "done", f"Indexed in {stored.get('collection') or 'Qdrant'}")
            yield sse({"type": "document", "document": stored, "journey_id": journey_id})
            yield sse({"type": "thinking", "text": "File stored in MongoDB and indexed."})
            yield sse({"type": "done", "document": stored, "journey_id": journey_id})
        except ValueError as exc:
            yield sse({"type": "error", "detail": str(exc)})
        except Exception as exc:
            yield sse({"type": "error", "detail": f"Could not ingest file: {exc}"})

    return StreamingResponse(
        generate(),
        media_type="text/event-stream",
        headers={
            "Cache-Control": "no-cache",
            "Connection": "keep-alive",
            "X-Accel-Buffering": "no",
        },
    )


@app.get("/documents")
def documents_list(journey_id: str = "", user: dict = Depends(current_user)):
    return {
        "documents": list_docs(user["user_id"], journey_id),
        "qdrant": qdrant_status(),
        "embeddings": embed_status(),
        "files": files_status(),
        "max_upload_mb": MAX_UPLOAD_MB,
    }


@app.get("/documents/{doc_id}/file")
def documents_download(doc_id: str, user: dict = Depends(current_user)):
    meta = get_doc(user["user_id"], doc_id)
    if not meta or not meta.get("gridfs_id"):
        raise HTTPException(status_code=404, detail="Original file is not stored")
    try:
        grid_out, file_meta = get_original_file(user["user_id"], meta["gridfs_id"])
    except FileNotFoundError as exc:
        raise HTTPException(status_code=404, detail=str(exc)) from exc
    except PermissionError as exc:
        raise HTTPException(status_code=403, detail=str(exc)) from exc
    filename = file_meta["filename"].replace('"', "")
    return StreamingResponse(
        grid_out,
        media_type=file_meta["content_type"],
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


@app.delete("/documents/{doc_id}")
def documents_delete(doc_id: str, user: dict = Depends(current_user)):
    return delete_doc(user["user_id"], doc_id)


@app.get("/journeys/{journey_id}")
def journeys_get(journey_id: str, user: dict = Depends(current_user)):
    journey = get_journey(user["user_id"], journey_id)
    loaded = load_all(user["user_id"], journey_id, journey.get("messages") or [], "")
    journey["messages"] = loaded["history"]
    journey["memory"] = {"layers": loaded["layers"], "facts": loaded["facts"]}
    return journey


@app.post("/chat/stream")
def chat_stream(req: ChatRequest, user: dict = Depends(current_user)):
    api_key = require_key()
    incoming = cleaned_messages(req)
    journey = resolve_journey(user, req.journey_id or req.session_id)
    journey_id = journey["journey_id"]
    user_id = user["user_id"]
    query = latest_user_text(incoming)
    loaded = load_all(user_id, journey_id, incoming, query)

    def generate():
        try:
            flow = []

            def step(name, status, detail=""):
                item = {"name": name, "status": status, "detail": detail}
                flow.append(item)
                return sse({"type": "flow", "steps": flow[:], "current": name})

            yield sse({"type": "thinking", "text": "Loading memory stores…"})
            yield step("memory", "done", "Loaded in-memory, Redis, thread, long-term")
            yield sse(
                {
                    "type": "memory",
                    "user_id": user_id,
                    "journey_id": journey_id,
                    "session_id": journey_id,
                    "layers": loaded["layers"],
                    "facts": loaded["facts"],
                }
            )

            yield sse({"type": "thinking", "text": "Searching uploaded documents in Qdrant…"})
            hits, rag_report = search_docs(user_id, query, journey_id)
            rag_notes = format_hits(hits)
            rag_key = hits_fingerprint(hits)
            yield step(
                "rag",
                rag_report.get("status") or "miss",
                rag_report.get("detail") or "Document search",
            )
            yield sse({"type": "retrieval", "report": rag_report, "hits": hits})

            yield sse({"type": "thinking", "text": "Checking prompt cache…"})
            cached, cache_report = get_prompt_cache(query, GROQ_MODEL, rag_key)
            yield sse({"type": "cache", "report": cache_report})

            if cached and cached.get("reply"):
                yield step("prompt_cache", "hit", cache_report.get("detail") or "Cache hit")
                analysis = cached.get("analysis") or DEFAULT_ANALYSIS
                yield sse({"type": "thinking", "text": "Using cached analysis…"})
                yield step("analyser", "cached", analysis.get("summary") or "Cached analysis")
                yield sse({"type": "analysis", "analysis": analysis, "model": GROQ_MODEL, "cached": True})
                yield sse({"type": "thinking", "text": "Replaying cached answer…"})
                yield step("generate", "cached", "Answer served from prompt cache")
                yield sse({"type": "token", "content": cached["reply"]})
                yield sse({"type": "done", "model": GROQ_MODEL, "cached": True})
                stored = loaded["history"] + [{"role": "assistant", "content": cached["reply"]}]
                writes = save_all(
                    journey_id,
                    user_id,
                    stored,
                    cached.get("title") or query,
                    cached["reply"],
                    analysis,
                )
                yield sse(
                    {
                        "type": "memory_write",
                        "writes": writes,
                        "journey_id": journey_id,
                        "session_id": journey_id,
                        "title": cached.get("title") or "",
                    }
                )
                if cached.get("followups"):
                    yield sse({"type": "followups", "questions": cached["followups"]})
                yield sse({"type": "thinking", "text": "Done (prompt cache)."})
                return

            yield step("prompt_cache", "miss", "No cached prompt, calling model")
            yield sse({"type": "thinking", "text": "Analysing the legal query…"})
            yield step("analyser", "running", "LangGraph analyse node")
            analysis = None
            reply_parts: list[str] = []
            for event in stream_graph(
                loaded["history"],
                api_key,
                GROQ_MODEL,
                memory_notes=loaded["notes"],
                rag_notes=rag_notes,
            ):
                if event.get("type") == "analysis":
                    analysis = event.get("analysis")
                    yield step("analyser", "done", (analysis or {}).get("summary") or "Analysed")
                    yield sse({"type": "thinking", "text": "Writing the answer…"})
                    yield step("generate", "running", "LangGraph generate node")
                if event.get("type") == "token":
                    reply_parts.append(event.get("content") or "")
                if event.get("type") == "done":
                    yield step("generate", "done", "Answer generated")
                yield sse(event)
            reply = "".join(reply_parts).strip()
            if reply:
                title = ""
                yield sse({"type": "thinking", "text": "Naming this chat…"})
                try:
                    title = suggest_title(query, reply, api_key, GROQ_MODEL)
                    yield step("title", "done", title)
                except Exception:
                    title = query[:60]
                    yield step("title", "done", title)
                followups: list[str] = []
                yield sse({"type": "thinking", "text": "Suggesting follow-up questions…"})
                try:
                    followups = suggest_followups(query, reply, api_key, GROQ_MODEL)
                    yield step("followups", "done", f"{len(followups)} questions")
                except Exception:
                    yield step("followups", "skip", "Could not suggest follow-ups")
                cache_write = set_prompt_cache(
                    query,
                    GROQ_MODEL,
                    {
                        "reply": reply,
                        "analysis": analysis or DEFAULT_ANALYSIS,
                        "followups": followups,
                        "title": title,
                    },
                    rag_key,
                )
                yield sse({"type": "cache_write", "report": cache_write})
                stored = loaded["history"] + [{"role": "assistant", "content": reply}]
                writes = save_all(journey_id, user_id, stored, title or query, reply, analysis)
                yield sse(
                    {
                        "type": "memory_write",
                        "writes": writes,
                        "journey_id": journey_id,
                        "session_id": journey_id,
                        "title": title,
                    }
                )
                if followups:
                    yield sse({"type": "followups", "questions": followups})
                yield sse({"type": "thinking", "text": "Done."})
        except Exception as exc:
            yield sse({"type": "error", "detail": f"LangGraph error: {exc}"})

    return StreamingResponse(
        generate(),
        media_type="text/event-stream",
        headers={
            "Cache-Control": "no-cache",
            "Connection": "keep-alive",
            "X-Accel-Buffering": "no",
        },
    )


@app.post("/chat")
def chat(req: ChatRequest, user: dict = Depends(current_user)):
    api_key = require_key()
    incoming = cleaned_messages(req)
    journey = resolve_journey(user, req.journey_id or req.session_id)
    journey_id = journey["journey_id"]
    user_id = user["user_id"]
    query = latest_user_text(incoming)
    loaded = load_all(user_id, journey_id, incoming, query)
    hits, rag_report = search_docs(user_id, query, journey_id)
    try:
        result = build_graph(api_key, GROQ_MODEL).invoke(
            {
                "messages": loaded["history"],
                "analysis": None,
                "reply": "",
                "memory_notes": loaded["notes"],
                "rag_notes": format_hits(hits),
            }
        )
    except Exception as exc:
        raise HTTPException(status_code=502, detail=f"LangGraph error: {exc}") from exc
    reply = (result.get("reply") or "").strip()
    if not reply:
        raise HTTPException(status_code=502, detail="Groq returned an empty reply")
    analysis = result.get("analysis") or DEFAULT_ANALYSIS
    stored = loaded["history"] + [{"role": "assistant", "content": reply}]
    writes = save_all(journey_id, user_id, stored, query, reply, analysis)
    return {
        "reply": reply,
        "model": GROQ_MODEL,
        "analysis": analysis,
        "journey_id": journey_id,
        "session_id": journey_id,
        "user_id": user_id,
        "memory": {"layers": loaded["layers"], "writes": writes, "facts": loaded["facts"]},
        "retrieval": {"report": rag_report, "hits": hits},
    }

# ═══════════════════════════════════════════════════════════════
# MULTI-AGENT CHAT ENDPOINTS
# ═══════════════════════════════════════════════════════════════

@app.post("/chat/stream/v2")
def chat_stream_v2(req: ChatRequest, user: dict = Depends(current_user)):
    """Multi-agent streaming chat — uses the root agent + specialist agents."""
    api_key = require_key()
    incoming = cleaned_messages(req)
    journey = resolve_journey(user, req.journey_id or req.session_id)
    journey_id = journey["journey_id"]
    user_id = user["user_id"]
    user_role = user.get("role") or ROLE_USER
    query = latest_user_text(incoming)
    loaded = load_all(user_id, journey_id, incoming, query)
    profile_text, _ = load_user_profile(user_id)
    
    # Context compression — summarize old messages if conversation is long
    compressed_history, compression_report = compress_history(
        loaded["history"], api_key, GROQ_MODEL, max_messages=8
    )
    if compression_report.get("used"):
        loaded["history"] = compressed_history

    def generate():
        try:
            flow: list[dict] = []
    
            def step(name, status, detail=""):
                """Emit a pipeline step (replaces any earlier step with the same name)."""
                item = {"name": name, "status": status, "detail": detail}
                flow[:] = [f for f in flow if f["name"] != name]
                flow.append(item)
                return sse({"type": "flow", "steps": flow[:], "current": name})
    
            # ── Greeting fast-path ──
            if _is_simple_greeting(query):
                reply = _GREETING_REPLY
                analysis = {
                    "intent": "other", "domain": "general", "complexity": "simple",
                    "jurisdiction": "unspecified", "on_topic": True,
                    "summary": "Greeting", "refined_query": query, "route_to": "assistant",
                }
                yield step("fast_path", "done", "Greeting detected — instant reply, no LLM call")
                yield sse({"type": "agent_route", "routed_to": "assistant", "analysis": analysis, "metadata": {}})
                yield sse({"type": "analysis", "analysis": analysis, "model": GROQ_MODEL})
                yield sse({"type": "token", "content": reply})
                stored = loaded["history"] + [{"role": "assistant", "content": reply}]
                writes = save_all(journey_id, user_id, stored, "Greeting", reply, analysis)
                yield step("memory_write", "done", f"{sum(1 for w in writes if w.get('wrote'))} memory store(s) updated")
                yield sse({"type": "memory_write", "writes": writes, "journey_id": journey_id, "title": "Greeting"})
                followups = [
                    "Can you help me draft a lease agreement?",
                    "What are the steps to file a small claims case?",
                    "What should I know about my tenant rights?",
                ]
                yield step("followups", "done", f"{len(followups)} canned suggestions (fast-path)")
                yield sse({"type": "followups", "questions": followups})
                yield sse({"type": "done", "model": GROQ_MODEL, "agent": "assistant"})
                return
    
            yield sse({"type": "thinking", "text": "Multi-agent pipeline starting…"})
            yield step("memory", "done", f"{len(loaded['layers'])} memory layer(s) · {len(loaded['facts'])} fact(s) loaded")
    
            if compression_report.get("used"):
                yield step(
                    "compress",
                    "done",
                    compression_report.get("detail") or "Old messages summarised to save tokens",
                )
    
            # RAG search
            yield sse({"type": "thinking", "text": "Searching uploaded documents (RAG)…"})
            hits, rag_report = search_docs(user_id, query, journey_id)
            rag_notes = format_hits(hits)
            yield step(
                "rag",
                "hit" if hits else "miss",
                f"{len(hits)} document chunk(s) matched in Qdrant" if hits else "No relevant documents found",
            )
            yield sse({"type": "retrieval", "report": rag_report, "hits": hits})
    
            # Stream through multi-agent graph
            reply_parts: list[str] = []
            analysis = None
            routed_to = "assistant"
    
            yield sse({"type": "thinking", "text": "Root agent classifying intent…"})
            yield step("orchestrator", "running", "Root agent reading query → intent · domain · complexity")
    
            for event in stream_multi_graph(
                loaded["history"],
                api_key,
                GROQ_MODEL,
                memory_notes=loaded["notes"],
                rag_notes=rag_notes,
                user_role=user_role,
                user_profile=profile_text,
                episodic_notes=loaded.get("episodic_notes", ""),
                procedural_notes=loaded.get("procedural_notes", ""),
            ):
                etype = event.get("type")
                if etype == "agent_route":
                    routed_to = event.get("routed_to") or "assistant"
                    analysis = event.get("analysis")
                    a = analysis or {}
                    yield step(
                        "orchestrator",
                        "done",
                        f"intent={a.get('intent')} · domain={a.get('domain')} · "
                        f"complexity={a.get('complexity')} → routed to {routed_to}",
                    )
                    yield sse({
                        "type": "agent_route",
                        "routed_to": routed_to,
                        "analysis": analysis,
                        "metadata": event.get("agent_metadata") or {},
                    })
                    yield sse({"type": "thinking", "text": f"Routed to {routed_to} agent…"})
                elif etype == "analysis":
                    analysis = event.get("analysis")
                    yield sse({"type": "analysis", "analysis": analysis, "model": GROQ_MODEL})
                elif etype == "agent_start":
                    agent = event.get("agent") or routed_to
                    yield step(agent, "running", f"{agent} agent generating reply…")
                    yield sse({"type": "thinking", "text": f"{agent} agent is writing…"})
                elif etype == "agent_done":
                    agent = event.get("agent") or routed_to
                    yield step(agent, "done", f"Reply generated · {event.get('reply_chars') or 0} characters")
                elif etype == "token":
                    reply_parts.append(event.get("content") or "")
                    yield sse(event)
                elif etype == "error":
                    yield sse(event)
                # "done" and unknown events are handled after the loop
    
            reply = "".join(reply_parts).strip()
            if reply:
                # Title + followups
                title = query[:60]
                yield step("title", "running", "LLM generating chat title…")
                try:
                    title = suggest_title(query, reply, api_key, GROQ_MODEL)
                    yield step("title", "done", f"“{title}”")
                except Exception:
                    yield step("title", "skip", "Using query excerpt as title")
                followups: list[str] = []
                yield step("followups", "running", "LLM generating follow-up questions…")
                try:
                    followups = suggest_followups(query, reply, api_key, GROQ_MODEL)
                    yield step("followups", "done", f"{len(followups)} question(s) generated from query + reply")
                except Exception:
                    yield step("followups", "skip", "Could not generate follow-ups")
    
                stored = loaded["history"] + [{"role": "assistant", "content": reply}]
                writes = save_all(journey_id, user_id, stored, title or query, reply, analysis)
                yield step(
                    "memory_write",
                    "done",
                    f"{sum(1 for w in writes if w.get('wrote'))} memory store(s) updated",
                )
                yield sse({
                    "type": "memory_write",
                    "writes": writes,
                    "journey_id": journey_id,
                    "title": title,
                })
                if followups:
                    yield sse({"type": "followups", "questions": followups})
                yield sse({"type": "done", "model": GROQ_MODEL, "agent": routed_to})
        except Exception as exc:
            yield sse({"type": "error", "detail": f"Multi-agent error: {exc}"})

    return StreamingResponse(
        generate(),
        media_type="text/event-stream",
        headers={"Cache-Control": "no-cache", "Connection": "keep-alive", "X-Accel-Buffering": "no"},
    )


@app.post("/chat/guest")
def chat_guest(req: ChatRequest, user: dict | None = Depends(optional_user)):
    """Guest mode chat — limited to 3 messages, no memory, no file upload."""
    api_key = require_key()
    incoming = cleaned_messages(req)
    if len(incoming) > 6:  # 3 user + 3 assistant = 6
        raise HTTPException(status_code=403, detail="Guest mode is limited to 3 messages. Please sign up for full access.")
    query = latest_user_text(incoming)

    # ── Greeting fast-path ──
    if _is_simple_greeting(query):
        return {
            "reply": _GREETING_REPLY,
            "model": GROQ_MODEL,
            "routed_to": "assistant",
            "analysis": {
                "intent": "other", "domain": "general", "complexity": "simple",
                "jurisdiction": "unspecified", "on_topic": True,
                "summary": "Greeting", "refined_query": query, "route_to": "assistant",
            },
            "guest_mode": True,
            "limit": "3 messages — sign up for unlimited access",
        }

    try:
        result_data = {}
        for event in stream_multi_graph(incoming, api_key, GROQ_MODEL, user_role="guest", user_profile=""):
            if event.get("type") == "token":
                result_data["reply"] = event.get("content") or ""
            if event.get("type") == "agent_route":
                result_data["routed_to"] = event.get("routed_to")
                result_data["analysis"] = event.get("analysis")
    except Exception as exc:
        raise HTTPException(status_code=502, detail=f"Error: {exc}") from exc

    reply = (result_data.get("reply") or "").strip()
    if not reply:
        raise HTTPException(status_code=502, detail="AI returned an empty reply")
    return {
        "reply": reply,
        "model": GROQ_MODEL,
        "routed_to": result_data.get("routed_to") or "assistant",
        "analysis": result_data.get("analysis") or {},
        "guest_mode": True,
        "limit": "3 messages — sign up for unlimited access",
    }


# ═══════════════════════════════════════════════════════════════
# USER PROFILE ENDPOINTS
# ═══════════════════════════════════════════════════════════════

class ProfileUpdateRequest(BaseModel):
    name: str = ""
    email: str = ""
    phone: str = ""
    facts: list[str] = []


@app.get("/memory/profile")
def memory_profile_get(user: dict = Depends(current_user)):
    """Get user profile."""
    profile = get_user_profile(user["user_id"])
    return {"profile": profile or {}}


@app.put("/memory/profile")
def memory_profile_put(req: ProfileUpdateRequest, user: dict = Depends(current_user)):
    """Update user profile."""
    updates = {}
    if req.name:
        updates["name"] = req.name
    if req.email:
        updates["email"] = req.email
    if req.phone:
        updates["phone"] = req.phone
    if req.facts:
        updates["facts"] = req.facts
    if not updates:
        raise HTTPException(status_code=400, detail="No fields to update")
    result = update_user_profile(user["user_id"], updates)
    if not result.get("ok"):
        raise HTTPException(status_code=500, detail=result.get("detail") or "Failed to update profile")
    return {"updated": result.get("updated"), "profile": get_user_profile(user["user_id"])}


@app.delete("/memory/profile")
def memory_profile_delete(user: dict = Depends(current_user)):
    """Delete user profile."""
    from memory import get_mongo, MONGO_DB, PROFILE_COLLECTION
    client = get_mongo()
    if client is None:
        raise HTTPException(status_code=500, detail="MongoDB unavailable")
    try:
        client[MONGO_DB][PROFILE_COLLECTION].delete_one({"user_id": user["user_id"]})
        return {"deleted": True}
    except Exception as exc:
        raise HTTPException(status_code=500, detail=str(exc))


# ═══════════════════════════════════════════════════════════════
# EPISODIC & PROCEDURAL MEMORY ENDPOINTS
# ═══════════════════════════════════════════════════════════════

@app.get("/memory/episodes")
def memory_episodes(user: dict = Depends(current_user)):
    """List all conversation episodes for this user."""
    return {"episodes": list_episodes(user["user_id"], limit=30)}


@app.get("/memory/preferences")
def memory_preferences(user: dict = Depends(current_user)):
    """Get user procedural preferences."""
    return {"preferences": get_procedural_memory(user["user_id"]) or {}}


# ═══════════════════════════════════════════════════════════════
# CONNECTOR ENDPOINTS
# ═══════════════════════════════════════════════════════════════

@app.get("/connectors")
def connectors_list():
    """List all available connectors and their status."""
    return {"connectors": list_connectors()}


@app.get("/connectors/{connector_name}")
def connector_detail(connector_name: str, query: str = ""):
    """Get details or search a specific connector."""
    conn = get_connector(connector_name)
    if not conn:
        raise HTTPException(status_code=404, detail=f"Connector '{connector_name}' not found")
    status = conn.status()
    if query and hasattr(conn, "search"):
        results = conn.search(query)
        return {"status": status, "search": results}
    return {"status": status}


# ═══════════════════════════════════════════════════════════════
# LAWYER FINDER & WEBSOCKET ENDPOINTS
# ═══════════════════════════════════════════════════════════════

class LawyerRoomCreate(BaseModel):
    lawyer_id: str
    journey_id: str = ""


@app.post("/lawyer/rooms")
def lawyer_room_create(payload: LawyerRoomCreate, user: dict = Depends(current_user)):
    """Create a real-time chat room between user and lawyer."""
    room = create_room(user["user_id"], payload.lawyer_id, payload.journey_id)
    return room


@app.get("/lawyer/rooms")
def lawyer_rooms_list(user: dict = Depends(current_user)):
    """List all chat rooms for the current user."""
    return {"rooms": list_rooms(user_id=user["user_id"])}


@app.get("/lawyer/rooms/{room_id}")
def lawyer_room_detail(room_id: str, user: dict = Depends(current_user)):
    """Get details of a specific chat room."""
    room = get_room(room_id)
    if not room:
        raise HTTPException(status_code=404, detail="Room not found")
    if room["user_id"] != user["user_id"] and room["lawyer_id"] != user["user_id"]:
        raise HTTPException(status_code=403, detail="Not authorized")
    return {
        "room_id": room["room_id"],
        "user_id": room["user_id"],
        "lawyer_id": room["lawyer_id"],
        "status": room["status"],
        "message_count": len(room["messages"]),
    }


@app.delete("/lawyer/rooms/{room_id}")
def lawyer_room_close(room_id: str, user: dict = Depends(current_user)):
    """Close a lawyer chat room."""
    room = get_room(room_id)
    if not room:
        raise HTTPException(status_code=404, detail="Room not found")
    if room["user_id"] != user["user_id"] and room["lawyer_id"] != user["user_id"]:
        raise HTTPException(status_code=403, detail="Not authorized")
    return close_room(room_id)


@app.websocket("/ws/lawyer/user/{room_id}")
async def ws_lawyer_user(websocket: WebSocket, room_id: str, user_id: str = ""):
    """WebSocket endpoint for user side of lawyer chat."""
    if not user_id:
        await websocket.close(code=4001, reason="user_id required")
        return
    await handle_user_websocket(websocket, room_id, user_id)


@app.websocket("/ws/lawyer/lawyer/{room_id}")
async def ws_lawyer_lawyer(websocket: WebSocket, room_id: str, lawyer_id: str = ""):
    """WebSocket endpoint for lawyer side of lawyer chat."""
    if not lawyer_id:
        await websocket.close(code=4001, reason="lawyer_id required")
        return
    await handle_lawyer_websocket(websocket, room_id, lawyer_id)


# ═══════════════════════════════════════════════════════════════
# ADMIN ENDPOINTS (role-restricted)
# ═══════════════════════════════════════════════════════════════

@app.get("/admin/system")
def admin_system_status(user: dict = Depends(require_role(ROLE_ADMIN))):
    """Full system status — admin only."""
    return {
        "health": {
            "provider": "groq",
            "model": GROQ_MODEL,
            "configured": bool(GROQ_API_KEY),
        },
        "agents": list_registered_agents(),
        "connectors": list_connectors(),
        "memory": layer_status(),
        "qdrant": qdrant_status(),
        "embeddings": embed_status(),
        "files": files_status(),
        "roles": ALL_ROLES,
    }


@app.get("/admin/agents")
def admin_list_agents(user: dict = Depends(require_role(ROLE_ADMIN))):
    """List all registered agents — admin only."""
    return {"agents": list_registered_agents()}


@app.get("/admin/connectors")
def admin_list_connectors(user: dict = Depends(require_role(ROLE_ADMIN))):
    """List all connectors — admin only."""
    return {"connectors": list_connectors()}


# ═══════════════════════════════════════════════════════════════
# LAWYER-SPECIFIC ENDPOINTS (role-restricted)
# ═══════════════════════════════════════════════════════════════

@app.get("/lawyer/my-rooms")
def lawyer_my_rooms(user: dict = Depends(require_role(ROLE_LAWYER, ROLE_ADMIN))):
    """List all rooms where this user is the lawyer."""
    return {"rooms": list_rooms(lawyer_id=user["user_id"])}


# ═══════════════════════════════════════════════════════════════
# EXPORT ENDPOINTS (PDF, DOCX, Image)
# ═══════════════════════════════════════════════════════════════

class ExportRequest(BaseModel):
    content: str
    format: str = "pdf"  # pdf | docx | txt
    title: str = "Legal Document"


class EmailSendRequest(BaseModel):
    to: list[str]
    subject: str
    body: str
    cc: list[str] = []
    bcc: list[str] = []


SMTP_HOST = os.getenv("SMTP_HOST", "smtp.gmail.com")
SMTP_PORT = int(os.getenv("SMTP_PORT", "587"))
SMTP_USER = os.getenv("SMTP_USER", "")
SMTP_PASS = os.getenv("SMTP_PASS", "")
SMTP_FROM = os.getenv("SMTP_FROM", SMTP_USER)


def _build_pdf(content: str, title: str) -> bytes:
    """Generate a PDF from text content using reportlab."""
    from io import BytesIO
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.lib.units import mm

    buf = BytesIO()
    doc = SimpleDocTemplate(
        buf, pagesize=A4,
        leftMargin=25*mm, rightMargin=25*mm,
        topMargin=20*mm, bottomMargin=20*mm,
    )
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "DocTitle", parent=styles["Heading1"],
        fontSize=16, spaceAfter=12, alignment=1,
    )
    body_style = ParagraphStyle(
        "DocBody", parent=styles["Normal"],
        fontSize=11, leading=15, spaceAfter=6,
    )
    bold_style = ParagraphStyle(
        "DocBold", parent=styles["Normal"],
        fontSize=11, leading=15, spaceAfter=6, fontName="Helvetica-Bold",
    )
    footer_style = ParagraphStyle(
        "DocFooter", parent=styles["Normal"],
        fontSize=8, textColor="grey", spaceBefore=20,
    )

    story = [Paragraph(title, title_style), Spacer(1, 6*mm)]

    for line in content.split("\n"):
        line = line.strip()
        if not line:
            story.append(Spacer(1, 3*mm))
        elif line.startswith("# "):
            story.append(Paragraph(line[2:], styles["Heading2"]))
        elif line.startswith("## "):
            story.append(Paragraph(line[3:], styles["Heading3"]))
        elif line.startswith("**") and line.endswith("**"):
            story.append(Paragraph(line.strip("*"), bold_style))
        elif line.startswith("- ") or line.startswith("* "):
            story.append(Paragraph(f"\u2022 {line[2:]}", body_style))
        else:
            # Escape XML for reportlab
            safe = line.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
            safe = safe.replace("**", "")
            story.append(Paragraph(safe, body_style))

    story.append(Spacer(1, 10*mm))
    story.append(Paragraph("Generated by Legal AI Assistant — Review before use.", footer_style))
    doc.build(story)
    return buf.getvalue()


def _build_docx(content: str, title: str) -> bytes:
    """Generate a DOCX from text content."""
    from io import BytesIO
    from docx import Document as DocxDocument
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = DocxDocument()

    # Title
    heading = doc.add_heading(title, level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Body
    for line in content.split("\n"):
        line = line.strip()
        if not line:
            doc.add_paragraph("")
        elif line.startswith("# "):
            doc.add_heading(line[2:], level=2)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=3)
        elif line.startswith("- ") or line.startswith("* "):
            doc.add_paragraph(line[2:], style="List Bullet")
        elif line.startswith("**") and line.endswith("**"):
            p = doc.add_paragraph()
            run = p.add_run(line.strip("*"))
            run.bold = True
        else:
            clean = line.replace("**", "")
            doc.add_paragraph(clean)

    # Footer
    doc.add_paragraph("")
    footer = doc.add_paragraph("Generated by Legal AI Assistant — Review before use.")
    footer.runs[0].font.size = Pt(8)
    footer.runs[0].font.italic = True

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


@app.post("/export/download")
def export_download(req: ExportRequest, user: dict = Depends(current_user)):
    """Export content as PDF, DOCX, or TXT for download."""
    from fastapi.responses import Response

    content = req.content.strip()
    title = req.title or "Legal Document"
    fmt = req.format.lower()

    if fmt == "pdf":
        try:
            data = _build_pdf(content, title)
        except Exception as exc:
            raise HTTPException(status_code=500, detail=f"PDF generation failed: {exc}")
        return Response(
            content=data,
            media_type="application/pdf",
            headers={"Content-Disposition": f'attachment; filename="{title.replace(" ", "_")}.pdf"'},
        )
    elif fmt == "docx":
        try:
            data = _build_docx(content, title)
        except Exception as exc:
            raise HTTPException(status_code=500, detail=f"DOCX generation failed: {exc}")
        return Response(
            content=data,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f'attachment; filename="{title.replace(" ", "_")}.docx"'},
        )
    elif fmt == "txt":
        data = f"{title}\n{'='*len(title)}\n\n{content}\n\n---\nGenerated by Legal AI Assistant"
        return Response(
            content=data.encode("utf-8"),
            media_type="text/plain",
            headers={"Content-Disposition": f'attachment; filename="{title.replace(" ", "_")}.txt"'},
        )
    else:
        raise HTTPException(status_code=400, detail=f"Unsupported format: {fmt}. Use pdf, docx, or txt.")


@app.post("/email/send")
def email_send(req: EmailSendRequest, user: dict = Depends(current_user)):
    """Send an email via SMTP."""
    import smtplib
    from email.mime.text import MIMEText
    from email.mime.multipart import MIMEMultipart

    if not SMTP_USER or not SMTP_PASS:
        raise HTTPException(status_code=503, detail="SMTP is not configured. Set SMTP_USER and SMTP_PASS env vars.")

    if not req.to:
        raise HTTPException(status_code=400, detail="At least one recipient is required.")

    msg = MIMEMultipart("alternative")
    msg["Subject"] = req.subject or "(No subject)"
    msg["From"] = SMTP_FROM or SMTP_USER
    msg["To"] = ", ".join(req.to)
    if req.cc:
        msg["Cc"] = ", ".join(req.cc)

    # Plain text body
    msg.attach(MIMEText(req.body, "plain", "utf-8"))

    # Also attach a simple HTML version
    html_body = req.body.replace("\n", "<br>")
    msg.attach(MIMEText(f"<html><body>{html_body}</body></html>", "html", "utf-8"))

    all_recipients = req.to + req.cc + req.bcc

    try:
        with smtplib.SMTP(SMTP_HOST, SMTP_PORT, timeout=30) as server:
            server.starttls()
            server.login(SMTP_USER, SMTP_PASS)
            server.sendmail(SMTP_FROM or SMTP_USER, all_recipients, msg.as_string())
    except Exception as exc:
        raise HTTPException(status_code=502, detail=f"Failed to send email: {exc}")

    return {
        "sent": True,
        "to": req.to,
        "cc": req.cc,
        "subject": req.subject,
        "detail": "Email sent successfully.",
    }
